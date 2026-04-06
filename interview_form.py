import customtkinter as ctk
import tkinter as tk
from tkinter import messagebox, filedialog, Toplevel
from datetime import datetime
from tkcalendar import Calendar
from database import get_connection
from PIL import Image, ImageTk
import os

# Word Export Libraries
from docx import Document
from docx.shared import Inches

# PDF Libraries
try:
    import pypdf
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
except ImportError:
    pypdf = None
    canvas = None
    letter = None

# Modern Colors
BG_COLOR = "white"
CARD_COLOR = "#F5F9FF"
TEXT_COLOR = "#333333"
ACCENT_COLOR = "#1966C7"
ASH = "#B0B0B0"


class InterviewForm:
    def __init__(self, parent, user_id, record_id=None):
        self.parent = parent
        self.user_id = user_id
        self.record_id = record_id
        self.image_path = None
        self.cv_path = None

        # Clear previous widgets
        for w in parent.winfo_children():
            w.destroy()

        # --- HEADER ---
        title_text = "Edit Interview" if self.record_id else "Interview Sheet"
        header = ctk.CTkFrame(parent, fg_color="white", corner_radius=0)
        header.pack(fill="x", pady=(10, 5))
        ctk.CTkLabel(header, text=title_text, font=("Arial", 24, "bold"), text_color="black").pack()

        # --- SCROLLABLE CONTAINER ---
        self.scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="white")
        self.scroll_frame.pack(fill="both", expand=True, padx=20, pady=10)

        # Form Card (Inner Container)
        self.form_area = ctk.CTkFrame(self.scroll_frame, fg_color=CARD_COLOR, corner_radius=15)
        self.form_area.pack(fill="both", expand=True, padx=10, pady=10)

        self.q_entries = {}
        self.row = 0

        # --- FORM BUILDER ---
        self.create_basic_details()
        self.create_timeline()
        self.create_location()
        self.create_project_section()
        self.create_questions_section()
        self.create_media_upload()
        self.create_action_buttons()
        self.create_summary_box()

        if self.record_id:
            self.load_data_for_edit()

    # --- UI SECTIONS ---
    def add_row(self, label_text, widget):
        ctk.CTkLabel(self.form_area, text=label_text, font=("Arial", 14), text_color=TEXT_COLOR, anchor="e").grid(
            row=self.row, column=0, sticky="e", padx=(20, 10), pady=10)
        widget.grid(row=self.row, column=1, sticky="w", padx=10, pady=10)
        self.row += 1

    def create_basic_details(self):
        # EP Name
        self.ep_entry = ctk.CTkEntry(self.form_area, width=300, placeholder_text="Enter EP Name")
        self.q_entries["EP Name"] = self.ep_entry
        self.add_row("EP Name:", self.ep_entry)

        # Gender
        gender_frame = ctk.CTkFrame(self.form_area, fg_color="transparent")
        self.gender_var = tk.StringVar()
        ctk.CTkRadioButton(gender_frame, text="Male", variable=self.gender_var, value="Male",
                           text_color=TEXT_COLOR).pack(side="left", padx=10)
        ctk.CTkRadioButton(gender_frame, text="Female", variable=self.gender_var, value="Female",
                           text_color=TEXT_COLOR).pack(side="left", padx=10)
        self.add_row("Gender:", gender_frame)

    def create_timeline(self):
        timeline_frame = ctk.CTkFrame(self.form_area, fg_color="transparent")

        self.start_date_entry = ctk.CTkEntry(timeline_frame, width=100, placeholder_text="Start")
        self.end_date_entry = ctk.CTkEntry(timeline_frame, width=100, placeholder_text="End")

        ctk.CTkButton(timeline_frame, text="📅", width=40,
                      command=lambda: self.open_calendar(self.start_date_entry)).pack(side="left", padx=5)
        self.start_date_entry.pack(side="left", padx=5)
        ctk.CTkLabel(timeline_frame, text="to", text_color=TEXT_COLOR).pack(side="left", padx=5)
        self.end_date_entry.pack(side="left", padx=5)
        ctk.CTkButton(timeline_frame, text="📅", width=40, command=lambda: self.open_calendar(self.end_date_entry)).pack(
            side="left", padx=5)

        self.add_row("Timeline:", timeline_frame)

    def create_location(self):
        self.lc_entry = ctk.CTkEntry(self.form_area, width=300, placeholder_text="Local Committee")
        self.q_entries["LC"] = self.lc_entry
        self.add_row("LC:", self.lc_entry)

        self.mc_entry = ctk.CTkEntry(self.form_area, width=300, placeholder_text="Member Committee")
        self.q_entries["MC"] = self.mc_entry
        self.add_row("MC:", self.mc_entry)

    def create_project_section(self):
        self.project_var = tk.StringVar()
        proj_frame = ctk.CTkFrame(self.form_area, fg_color="transparent")
        self.add_row("Project:", proj_frame)

        projects = ["Aquatica", "Green Leaders", "Global Classroom", "On The Map", "Rooted"]
        self.projects = projects

        # Grid layout for radio buttons
        for i, p in enumerate(projects):
            ctk.CTkRadioButton(proj_frame, text=p, variable=self.project_var, value=p,
                               command=self.load_project_questions, text_color=TEXT_COLOR).grid(row=0, column=i, padx=5)

    def create_questions_section(self):
        self.common_questions = [
            "Introduction", "Unique Qualities", "Previous volunteering experience",
            "Get to know about AIESEC", "Reason for choosing Sri Lanka"
        ]
        self.project_questions = {
            "Aquatica": ["Why improve sustainability?", "Marine ecosystem contributions?"],
            "Green Leaders": ["Climate change importance?", "Contribution plans?"],
            "Global Classroom": ["Quality education importance?", "Handling students?"],
            "On The Map": ["Skills for volunteering?", "Travel difficulties?"],
            "Rooted": ["Pet care experience?", "Fieldwork willingness?"]
        }

        # Dropdown
        self.add_row("Add Question:", ctk.CTkLabel(self.form_area, text="Select below to add:", text_color="gray"))

        self.dropdown_var = tk.StringVar()
        self.dropdown = ctk.CTkComboBox(self.form_area, variable=self.dropdown_var, values=self.common_questions,
                                        width=300)
        self.dropdown.grid(row=self.row, column=1, sticky="w", padx=10)

        ctk.CTkButton(self.form_area, text="+ Add", width=60, command=self.add_optional_entry).grid(row=self.row,
                                                                                                    column=1,
                                                                                                    sticky="e", padx=50)
        self.row += 1

        # Container for added questions
        self.optional_container = ctk.CTkFrame(self.form_area, fg_color="transparent")
        self.optional_container.grid(row=self.row, column=0, columnspan=2, sticky="w", padx=20)
        self.selected_optional_entries = {}
        self.row += 1

        # Text Areas
        self.remarks_box = ctk.CTkTextbox(self.form_area, height=80, width=500, border_color=ASH, border_width=2)
        self.add_row("Remarks:", self.remarks_box)

        self.ep_questions_box = ctk.CTkTextbox(self.form_area, height=80, width=500, border_color=ASH, border_width=2)
        self.add_row("EP Questions:", self.ep_questions_box)

    def create_media_upload(self):
        media_frame = ctk.CTkFrame(self.form_area, fg_color="transparent")
        media_frame.grid(row=self.row, column=1, sticky="w", pady=10)

        ctk.CTkButton(media_frame, text="Upload Photo", command=self.upload_image, fg_color="#E0E0E0",
                      text_color="black").pack(side="left", padx=5)
        self.img_label_text = ctk.CTkLabel(media_frame, text="No image", text_color="gray")
        self.img_label_text.pack(side="left", padx=5)

        ctk.CTkButton(media_frame, text="Attach CV", command=self.upload_cv, fg_color="#E0E0E0",
                      text_color="black").pack(side="left", padx=(20, 5))
        self.cv_label = ctk.CTkLabel(media_frame, text="No CV", text_color="gray")
        self.cv_label.pack(side="left", padx=5)

        self.row += 1

        # Image Preview Label
        self.img_preview_label = ctk.CTkLabel(self.form_area, text="")
        self.img_preview_label.grid(row=self.row, column=1, sticky="w")
        self.row += 1

    def create_action_buttons(self):
        btn_frame = ctk.CTkFrame(self.scroll_frame, fg_color="transparent")
        btn_frame.pack(fill="x", pady=20)

        main_btn_text = "UPDATE RECORD" if self.record_id else "SAVE RECORD"

        # Main Save
        ctk.CTkButton(btn_frame, text=main_btn_text, command=self.enter_data,
                      fg_color=ACCENT_COLOR, height=40).pack(side="left", padx=5, expand=True)

        # Exports
        ctk.CTkButton(btn_frame, text="Export PDF", command=self.export_pdf,
                      fg_color="#D32F2F", height=40).pack(side="left", padx=5, expand=True)
        ctk.CTkButton(btn_frame, text="Export Word", command=self.export_word,
                      fg_color="#388E3C", height=40).pack(side="left", padx=5, expand=True)

        # Navigation
        ctk.CTkButton(btn_frame, text="Close", command=self.save_and_close,
                      fg_color="gray", height=40).pack(side="left", padx=5, expand=True)

    def create_summary_box(self):
        ctk.CTkLabel(self.scroll_frame, text="Summary Preview", font=("Arial", 16, "bold")).pack(anchor="w", padx=10)
        self.summary_box = ctk.CTkTextbox(self.scroll_frame, height=150, border_color=ASH, border_width=2)
        self.summary_box.pack(fill="x", padx=10, pady=(5, 20))

    # --- LOGIC METHODS ---

    def load_project_questions(self):
        sel = self.project_var.get()
        if not sel: return
        pq = self.project_questions.get(sel, [])
        self.dropdown.configure(values=self.common_questions + pq)
        self.dropdown.set("")

    def add_optional_entry(self, question=None, answer=""):
        q = question or self.dropdown.get()
        if not q or q in self.selected_optional_entries: return

        f = ctk.CTkFrame(self.optional_container, fg_color="transparent")
        f.pack(fill="x", pady=2)

        ctk.CTkLabel(f, text=q, width=150, anchor="w", text_color=TEXT_COLOR).pack(side="left")

        entry = ctk.CTkEntry(f, width=300)
        if answer: entry.insert(0, answer)
        entry.pack(side="left", padx=5)

        ctk.CTkButton(f, text="x", width=30, fg_color="#FF5252", command=lambda: self.remove_optional(q, f)).pack(
            side="left")

        self.selected_optional_entries[q] = entry

    def remove_optional(self, q, f):
        del self.selected_optional_entries[q]
        f.destroy()

    def open_calendar(self, target_entry):
        top = Toplevel(self.parent)
        cal = Calendar(top, selectmode="day", date_pattern="yyyy-mm-dd")
        cal.pack(fill="both", expand=True)

        def pick():
            target_entry.delete(0, tk.END)
            target_entry.insert(0, cal.get_date())
            top.destroy()

        ctk.CTkButton(top, text="Select", command=pick).pack(pady=5)

    def upload_image(self):
        fp = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg")])
        if fp:
            self.image_path = fp
            self.img_label_text.configure(text="Image Loaded", text_color="green")
            #self.display_image(fp)

    def display_image(self, fp):
        if os.path.exists(fp):
            img = ctk.CTkImage(light_image=Image.open(fp), size=(150, 150))
            self.img_preview_label.configure(image=img, text="")

    def upload_cv(self):
        fp = filedialog.askopenfilename(filetypes=[("Documents", "*.docx *.pdf")])
        if fp:
            self.cv_path = fp
            self.cv_label.configure(text="CV Attached", text_color="green")

    def enter_data(self):
        ep_name = self.ep_entry.get().strip()
        project = self.project_var.get()

        if not ep_name or not project:
            messagebox.showerror("Error", "Name and Project are required.")
            return

        # 1. Build the Visible Summary (Clean, no paths)
        summary_visible = f"EP Name: {ep_name}\n"
        summary_visible += f"Gender: {self.gender_var.get()}\n"
        summary_visible += f"Timeline: {self.start_date_entry.get()} -> {self.end_date_entry.get()}\n"
        summary_visible += f"LC: {self.lc_entry.get()}\nMC: {self.mc_entry.get()}\n"
        summary_visible += f"Project: {project}\n"

        for q, e in self.selected_optional_entries.items():
            if e.get(): summary_visible += f"{q}: {e.get()}\n"

        summary_visible += f"Remarks: {self.remarks_box.get('1.0', 'end').strip()}\n"
        summary_visible += f"EP Questions: {self.ep_questions_box.get('1.0', 'end').strip()}\n"
        summary_visible += f"Entered/Updated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

        # 2. Build the Database Summary (Must include paths for persistence)
        summary_db = summary_visible
        if self.image_path: summary_db += f"\nImage Path: {self.image_path}"
        if self.cv_path: summary_db += f"\nCV Path: {self.cv_path}"

        conn, cursor = get_connection()
        try:
            if self.record_id:
                cursor.execute("UPDATE interviews SET summary=%s WHERE id=%s", (summary_db, self.record_id))
            else:
                cursor.execute("INSERT INTO interviews (user_id, summary) VALUES (%s, %s)", (self.user_id, summary_db))
            conn.commit()

            # Update UI with the CLEAN version
            self.summary_box.delete("1.0", "end")
            self.summary_box.insert("end", summary_visible)
            messagebox.showinfo("Success", "Record Saved Successfully!")
        except Exception as e:
            messagebox.showerror("Error", str(e))
        finally:
            cursor.close()
            conn.close()

    def load_data_for_edit(self):
        conn, cursor = get_connection()
        try:
            cursor.execute("SELECT summary FROM interviews WHERE id=%s", (self.record_id,))
            res = cursor.fetchone()
            if not res: return

            raw = res[0]

            # Simple Parsing
            data = {}
            for line in raw.split("\n"):
                if ":" in line:
                    k, v = line.split(":", 1)
                    data[k.strip()] = v.strip()

            if "EP Name" in data: self.ep_entry.insert(0, data["EP Name"])
            if "LC" in data: self.lc_entry.insert(0, data["LC"])
            if "MC" in data: self.mc_entry.insert(0, data["MC"])
            if "Gender" in data: self.gender_var.set(data["Gender"])
            if "Project" in data:
                self.project_var.set(data["Project"])
                self.load_project_questions()

            if "Timeline" in data:
                parts = data["Timeline"].split("->")
                if len(parts) == 2:
                    self.start_date_entry.insert(0, parts[0].strip())
                    self.end_date_entry.insert(0, parts[1].strip())

            if "Remarks" in data: self.remarks_box.insert("1.0", data["Remarks"])
            if "EP Questions" in data: self.ep_questions_box.insert("1.0", data["EP Questions"])

            # Extract paths silently
            if "Image Path" in data:
                self.image_path = data["Image Path"]
                self.img_label_text.configure(text="Image Loaded", text_color="green")
                #self.display_image(self.image_path)
            if "CV Path" in data:
                self.cv_path = data["CV Path"]
                self.cv_label.configure(text="CV Attached", text_color="green")

            # Fill optional entries
            std_keys = ["EP Name", "Gender", "Timeline", "LC", "MC", "Project", "Remarks", "EP Questions",
                        "Image Path", "CV Path", "Entered/Updated at"]
            for k, v in data.items():
                if k not in std_keys and k:
                    self.add_optional_entry(k, v)

            # Show CLEAN version in the summary box (filter out path lines)
            clean_lines = [line for line in raw.split('\n')
                           if not line.startswith("Image Path:") and not line.startswith("CV Path:")]
            self.summary_box.insert("end", "\n".join(clean_lines))

        finally:
            cursor.close()
            conn.close()

    # --- EXPORT WORD (Fixed) ---
    def export_word(self):
        if not self.summary_box.get("1.0", "end").strip():
            messagebox.showerror("Error", "No summary to export")
            return

        doc = Document()
        doc.add_heading("Interview Summary", 0)

        # 1. Place Image at Top
        if self.image_path and os.path.exists(self.image_path):
            try:
                doc.add_picture(self.image_path, width=Inches(3))
            except Exception as e:
                doc.add_paragraph(f"[Could not load image: {e}]")

        # 2. Add Summary Text (This is already clean, as summary_box doesn't have paths)
        doc.add_paragraph(self.summary_box.get("1.0", "end"))

        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if save_path:
            try:
                doc.save(save_path)
                #messagebox.showinfo("Saved", f"Summary exported to {save_path}")
                messagebox.showinfo("Saved", " Document Exported Successfully")

            except Exception as e:
                messagebox.showerror("Error", f"Failed to save Word doc: {e}")

    # --- EXPORT PDF (Fixed: Image at Top & Larger) ---
    def export_pdf(self):
        if not self.summary_box.get("1.0", "end").strip():
            messagebox.showerror("Error", "No summary to export")
            return

        path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
        if not path:
            return

        if not canvas:
            messagebox.showerror("Error", "ReportLab library not installed. Cannot export PDF.")
            return

        try:
            # Create temporary PDF for summary
            temp_pdf = "temp_summary.pdf"
            c = canvas.Canvas(temp_pdf, pagesize=letter)
            page_width, page_height = letter
            y = page_height - 50  # Start cursor near top

            # 1. HEADER TITLE
            c.setFont("Helvetica-Bold", 18)
            c.drawString(50, y, "Interview Summary")
            y -= 40  # Move down for image

            # 2. DRAW IMAGE (Dynamic Height Calculation)
            if self.image_path and os.path.exists(self.image_path):
                try:
                    # Open image to get dimensions
                    pil_img = Image.open(self.image_path)
                    img_w, img_h = pil_img.size

                    # Define max dimensions
                    max_w, max_h = 250, 250

                    # Calculate aspect ratio to fit within max_w/max_h
                    ratio = min(max_w / img_w, max_h / img_h)

                    # Calculate new actual dimensions
                    draw_w = img_w * ratio
                    draw_h = img_h * ratio

                    # Draw image at the calculated height
                    # Note: y is the bottom-left corner, so we subtract draw_h
                    c.drawImage(self.image_path, 50, y - draw_h, width=draw_w, height=draw_h,
                                preserveAspectRatio=True, mask='auto')

                    # Move cursor down by the ACTUAL height used + small padding (20)
                    y -= (draw_h + 20)

                except Exception as e:
                    print(f"Image draw failed: {e}")
                    c.drawString(50, y, "[Image Load Failed]")
                    y -= 20

            # 3. DRAW TEXT
            c.setFont("Helvetica", 12)
            # Get text and remove empty leading/trailing lines
            text = self.summary_box.get("1.0", "end").strip()

            for line in text.split("\n"):
                # Check if we need a new page
                if y < 50:
                    c.showPage()
                    y = page_height - 50
                    c.setFont("Helvetica", 12)  # Reset font on new page

                c.drawString(50, y, line)
                y -= 18  # Line spacing

            c.save()

            # 4. MERGE WITH ATTACHED CV (If exists)
            writer = pypdf.PdfWriter()
            writer.append(temp_pdf)

            if self.cv_path and self.cv_path.endswith(".pdf") and os.path.exists(self.cv_path):
                try:
                    writer.append(self.cv_path)
                except Exception as e:
                    messagebox.showwarning("Warning", f"Could not append CV: {e}")

            with open(path, "wb") as f:
                writer.write(f)

            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)

            messagebox.showinfo("Success", "PDF Exported Successfully")

        except Exception as e:
            messagebox.showerror("Error", f"PDF Export failed: {e}")

    def save_and_close(self):
        # Use winfo_toplevel() to find the main application window directly
        main_window = self.parent.winfo_toplevel()

        # Try to find the controller on the main window
        if hasattr(main_window, "controller"):
            controller = main_window.controller
        elif hasattr(main_window, "show_frame"):
            # The main window itself might be the controller
            controller = main_window
        else:
            # Fallback for deep nesting or different structure
            try:
                controller = self.parent.master.master.controller
            except AttributeError:
                messagebox.showerror("Error", "Could not find application controller.")
                return

        try:
            user = controller.current_user
            # Check admin status (index 8 is commonly 'is_admin')
            is_admin = user[8] if user and len(user) > 8 else 0

            if is_admin:
                controller.show_frame("AdminInterviewerModeFrame")
            else:
                controller.show_frame("DashboardFrame")
        except Exception as e:
            messagebox.showerror("Error", f"Navigation failed: {e}")